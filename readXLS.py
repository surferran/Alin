"""
by:
 https://stackoverflow.com/questions/40596518/writing-a-python-pandas-dataframe-to-word-document

https://python-docx.readthedocs.io/en/latest/

"""

import pandas as pd

fileName=r'תשלום דיבידנד לחברים 2019.xlsx'

df = pd.read_excel(fileName)
df['Unnamed: 3']
dfTrimmed = df['Unnamed: 3'].dropna()
#print (df)

import docx
from docx import Document

def op1():
    from docx import Document
    from docx.shared import Inches
    
    document = Document()
    
    document.add_heading('Document Title', 0)
    
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')
    
    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    
#    document.add_picture('monty-truth.png', width=Inches(1.25))
    
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )
    
#    table = document.add_table(rows=1, cols=3)
    table = document.add_table(rows=df['Unnamed: 3'].size, cols=1)
    
    hdr_cells = table.rows[0].cells
#    hdr_cells[0].text = 'Qty'
#    hdr_cells[1].text = 'Id'
    hdr_cells[0].text = 'Desc'
#    for qty, id, desc in records:
    aLst = df['Unnamed: 3'].dropna().values
    for desc in aLst:
        
        row_cells = table.add_row().cells
#        row_cells[0].text = str(qty)
#        row_cells[1].text = id
        row_cells[0].text = desc
        
        row_cells = table.add_row().cells
        row_cells[0].text = desc
    
#        document.add_page_break()
        
    document.add_page_break()
    
    document.save('demo.docx')

#op1()

CHOOSE_OPTION=1
if CHOOSE_OPTION==1:

    # open an existing document
#    doc = docx.Document('test.docx')
    
    doc = Document()
    
    
    df = dfTrimmed
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
#    t = doc.add_table(df.shape[0]+1, df.shape[1])
    t = doc.add_table(df.shape[0]+1, 1)
    
    # add the header rows.                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                               
    
    # add the rest of the data frame
    for i in range(df.shape[0]):
#        for j in range(df.shape[-1]):
        j=0
#        t.cell(i+1,j).text = str(df.values[i,j])
        t.cell(i+1,0).text = str(df.values[i])
    
    # save the doc
    doc.save('test.docx')
    
elif CHOOSE_OPTION==2:
    import win32
    import pywin as win32
    wordApp = win32.gencache.EnsureDispatch('Word.Application')
    
    wordApp.Visible = False
    doc = wordApp.Documents.Open(os.getcwd()+'\\temp.docx')
    rng = doc.Bookmarks("PUTTABLEHERE").Range
    rng.InsertTable.here